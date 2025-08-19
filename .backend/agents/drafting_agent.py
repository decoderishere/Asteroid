import os
import json
from typing import Dict, Any, List
from sqlalchemy.orm import Session
from .base_agent import BaseAgent
from schemas import AgentInput, AgentOutput
from models import Document, PermittingTemplate, ProjectFile
from docx import Document as DocxDocument
from docx.shared import Inches
import uuid

class DraftingAgent(BaseAgent):
    def __init__(self, openrouter_client):
        super().__init__("DraftingAgent", openrouter_client)
        self.chile_document_types = [
            "environmental_impact_assessment",
            "interconnection_request",
            "land_use_permit",
            "construction_permit",
            "electrical_safety_certification",
            "environmental_compliance_report",
            "substation_connection_study",
            "grid_impact_analysis"
        ]
    
    def process(self, input_data: AgentInput, db: Session) -> AgentOutput:
        task_type = input_data.task_type
        project_id = input_data.project_id
        
        if task_type == "generate_initial_drafts":
            return self._generate_initial_drafts(input_data, db)
        elif task_type == "update_affected_documents":
            return self._update_affected_documents(input_data, db)
        elif task_type == "generate_single_document":
            return self._generate_single_document(input_data, db)
        else:
            return AgentOutput(
                agent_name=self.name,
                task_type=task_type,
                project_id=project_id,
                output_data={},
                model_used="none",
                reasoning=f"Unknown drafting task: {task_type}",
                execution_time=0.0,
                success=False,
                error_message=f"Unknown task type: {task_type}"
            )
    
    def _generate_initial_drafts(self, input_data: AgentInput, db: Session) -> AgentOutput:
        project_id = input_data.project_id
        extracted_data_list = input_data.input_data.get("extracted_data", [])
        
        # Combine all extracted data
        combined_data = self._combine_extracted_data(extracted_data_list)
        
        generated_docs = []
        
        for doc_type in self.chile_document_types:
            if self._should_generate_document(doc_type, combined_data):
                doc_result = self._create_document(project_id, doc_type, combined_data, db)
                if doc_result:
                    generated_docs.append(doc_result)
        
        return AgentOutput(
            agent_name=self.name,
            task_type=input_data.task_type,
            project_id=project_id,
            output_data={
                "generated_documents": generated_docs,
                "document_count": len(generated_docs)
            },
            model_used=self.openrouter_client.get_optimal_model("drafting"),
            reasoning=f"Generated {len(generated_docs)} initial draft documents based on available data",
            execution_time=0.0,
            success=True
        )
    
    def _update_affected_documents(self, input_data: AgentInput, db: Session) -> AgentOutput:
        project_id = input_data.project_id
        new_data = input_data.input_data.get("new_extracted_data", {})
        
        # Find existing documents that need updates
        existing_docs = db.query(Document).filter(
            Document.project_id == project_id,
            Document.status.in_(["draft", "needs_review"])
        ).all()
        
        updated_docs = []
        
        for doc in existing_docs:
            if self._document_needs_update(doc, new_data):
                updated_doc = self._update_document(doc, new_data, db)
                if updated_doc:
                    updated_docs.append(updated_doc)
        
        return AgentOutput(
            agent_name=self.name,
            task_type=input_data.task_type,
            project_id=project_id,
            output_data={
                "updated_documents": updated_docs,
                "update_count": len(updated_docs)
            },
            model_used=self.openrouter_client.get_optimal_model("drafting"),
            reasoning=f"Updated {len(updated_docs)} documents with new information",
            execution_time=0.0,
            success=True
        )
    
    def _generate_single_document(self, input_data: AgentInput, db: Session) -> AgentOutput:
        project_id = input_data.project_id
        document_type = input_data.input_data.get("document_type")
        project_data = input_data.input_data.get("project_data", {})
        
        if not document_type:
            return AgentOutput(
                agent_name=self.name,
                task_type=input_data.task_type,
                project_id=project_id,
                output_data={},
                model_used="none",
                reasoning="Document type not specified",
                execution_time=0.0,
                success=False,
                error_message="document_type is required"
            )
        
        # Create combined data from project information
        combined_data = {
            "project_details": {
                "name": project_data.get("name"),
                "substation_id": project_data.get("substation_id"),
                "substation_name": project_data.get("substation_name"),
                "project_developer": project_data.get("project_developer"),
                "description": project_data.get("description")
            },
            "technical_specs": {
                "capacity_mw": project_data.get("capacity_mw"),
                "voltage_level": project_data.get("voltage_level"),
                "technology_type": project_data.get("technology_type"),
                "grid_connection_type": project_data.get("grid_connection_type")
            },
            "environmental_data": {
                "latitude": project_data.get("latitude"),
                "longitude": project_data.get("longitude")
            },
            "regulatory_info": {},
            "timeline": {},
            "contacts": {},
            "source_files": []
        }
        
        # Generate the document
        doc_result = self._create_document(project_id, document_type, combined_data, db)
        
        if not doc_result:
            return AgentOutput(
                agent_name=self.name,
                task_type=input_data.task_type,
                project_id=project_id,
                output_data={},
                model_used=self.openrouter_client.get_optimal_model("drafting"),
                reasoning=f"Failed to generate {document_type} document",
                execution_time=0.0,
                success=False,
                error_message=f"Document generation failed for type: {document_type}"
            )
        
        return AgentOutput(
            agent_name=self.name,
            task_type=input_data.task_type,
            project_id=project_id,
            output_data={
                "document_id": doc_result["document_id"],
                "doc_type": document_type,
                "title": doc_result["title"],
                "placeholders_count": doc_result.get("placeholders_count", 0)
            },
            model_used=self.openrouter_client.get_optimal_model("drafting"),
            reasoning=f"Successfully generated {document_type} document",
            execution_time=0.0,
            success=True
        )
    
    def _combine_extracted_data(self, extracted_data_list: List[Dict]) -> Dict[str, Any]:
        combined = {
            "project_details": {},
            "technical_specs": {},
            "environmental_data": {},
            "regulatory_info": {},
            "timeline": {},
            "contacts": {},
            "source_files": []
        }
        
        for data in extracted_data_list:
            if isinstance(data, dict) and "extracted_data" in data:
                extracted = data["extracted_data"]
                combined["source_files"].append(data.get("file_id", "unknown"))
                
                # Merge data intelligently
                for key in combined.keys():
                    if key in extracted and isinstance(extracted[key], dict):
                        combined[key].update(extracted[key])
                    elif key in extracted:
                        combined[key] = extracted[key]
        
        return combined
    
    def _should_generate_document(self, doc_type: str, data: Dict[str, Any]) -> bool:
        # Basic heuristics for document generation
        requirements = {
            "environmental_impact_assessment": ["environmental_data", "project_details"],
            "interconnection_request": ["technical_specs", "project_details"],
            "land_use_permit": ["project_details", "environmental_data"],
            "construction_permit": ["project_details", "technical_specs"],
            "electrical_safety_certification": ["technical_specs"],
            "environmental_compliance_report": ["environmental_data"],
            "substation_connection_study": ["technical_specs"],
            "grid_impact_analysis": ["technical_specs"]
        }
        
        required_keys = requirements.get(doc_type, ["project_details"])
        return any(data.get(key) for key in required_keys)
    
    def _create_document(self, project_id: str, doc_type: str, data: Dict[str, Any], db: Session) -> Dict[str, Any]:
        # Get template if available
        template = self._get_chile_template(doc_type)
        
        # Generate content using LLM
        content = self._generate_document_content(doc_type, data, template)
        
        if not content:
            return None
        
        # Create document record
        document = Document(
            project_id=project_id,
            doc_type=doc_type,
            title=self._get_document_title(doc_type, data),
            content=content["text"],
            source_files=data.get("source_files", []),
            placeholders=content.get("placeholders", [])
        )
        
        db.add(document)
        db.commit()
        
        # Generate DOCX file
        file_path = self._create_docx_file(document.id, doc_type, content, project_id)
        document.file_path = file_path
        db.commit()
        
        return {
            "document_id": document.id,
            "doc_type": doc_type,
            "title": document.title,
            "file_path": file_path,
            "placeholders_count": len(content.get("placeholders", []))
        }
    
    def _get_chile_template(self, doc_type: str) -> str:
        # Chile-specific document templates
        templates = {
            "environmental_impact_assessment": """
            ESTUDIO DE IMPACTO AMBIENTAL
            PROYECTO SISTEMA DE ALMACENAMIENTO DE ENERGÍA EN BATERÍAS (BESS)
            
            1. DESCRIPCIÓN DEL PROYECTO
            {project_description}
            
            2. LOCALIZACIÓN
            {project_location}
            
            3. CARACTERÍSTICAS TÉCNICAS
            {technical_specifications}
            
            4. ANÁLISIS DE IMPACTO AMBIENTAL
            {environmental_impact}
            
            5. MEDIDAS DE MITIGACIÓN
            {mitigation_measures}
            
            6. PLAN DE SEGUIMIENTO
            {monitoring_plan}
            """,
            
            "interconnection_request": """
            SOLICITUD DE INTERCONEXIÓN AL SISTEMA ELÉCTRICO NACIONAL
            
            1. DATOS DEL SOLICITANTE
            {applicant_data}
            
            2. CARACTERÍSTICAS DEL PROYECTO
            {project_characteristics}
            
            3. PUNTO DE CONEXIÓN SOLICITADO
            {connection_point}
            
            4. ESPECIFICACIONES TÉCNICAS DEL EQUIPAMIENTO
            {equipment_specifications}
            
            5. ESTUDIOS DE CONEXIÓN
            {connection_studies}
            
            6. CRONOGRAMA DE IMPLEMENTACIÓN
            {implementation_schedule}
            """
        }
        
        return templates.get(doc_type, "Documento tipo {doc_type} - Contenido a generar.")
    
    def _generate_document_content(self, doc_type: str, data: Dict[str, Any], template: str) -> Dict[str, Any]:
        system_message = self.openrouter_client.create_system_message(
            f"""You are an expert in Chilean BESS permitting documentation. Generate a complete {doc_type} document 
            following Chilean regulations and best practices.
            
            Use the provided template structure and fill it with information from the extracted data.
            Flag any missing critical information as placeholders.
            
            Requirements:
            - Use formal Spanish technical language appropriate for Chilean regulatory submissions
            - Include all mandatory sections per Chilean regulations
            - Cite specific Chilean laws/regulations where applicable (Ley General de Servicios Eléctricos, etc.)
            - Flag missing information clearly with [PLACEHOLDER: description] format
            - Ensure technical accuracy for BESS systems
            
            Return JSON with:
            - "text": Complete document content
            - "placeholders": List of missing information items
            - "regulatory_references": Relevant Chilean regulations cited
            """
        )
        
        user_message = self.openrouter_client.create_user_message(
            f"""Document Type: {doc_type}
            
            Template:
            {template}
            
            Available Data:
            {json.dumps(data, indent=2, ensure_ascii=False)}
            
            Generate the complete document content."""
        )
        
        response = self.openrouter_client.chat_completion(
            messages=[system_message, user_message],
            task_type="drafting",
            temperature=0.4
        )
        
        if response["success"]:
            try:
                content = response["data"]["choices"][0]["message"]["content"]
                if "```json" in content:
                    json_start = content.find("```json") + 7
                    json_end = content.find("```", json_start)
                    json_str = content[json_start:json_end].strip()
                else:
                    json_str = content.strip()
                
                return json.loads(json_str)
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                return {
                    "text": response["data"]["choices"][0]["message"]["content"],
                    "placeholders": ["JSON parsing failed - manual review needed"],
                    "regulatory_references": []
                }
        
        return None
    
    def _get_document_title(self, doc_type: str, data: Dict[str, Any]) -> str:
        project_name = data.get("project_details", {}).get("name", "Proyecto BESS")
        
        titles = {
            "environmental_impact_assessment": f"Estudio de Impacto Ambiental - {project_name}",
            "interconnection_request": f"Solicitud de Interconexión - {project_name}",
            "land_use_permit": f"Permiso de Uso de Suelo - {project_name}",
            "construction_permit": f"Permiso de Construcción - {project_name}",
            "electrical_safety_certification": f"Certificación de Seguridad Eléctrica - {project_name}",
            "environmental_compliance_report": f"Reporte de Cumplimiento Ambiental - {project_name}",
            "substation_connection_study": f"Estudio de Conexión a Subestación - {project_name}",
            "grid_impact_analysis": f"Análisis de Impacto en la Red - {project_name}"
        }
        
        return titles.get(doc_type, f"Documento {doc_type} - {project_name}")
    
    def _create_docx_file(self, doc_id: str, doc_type: str, content: Dict[str, Any], project_id: str) -> str:
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(content.get("title", f"Documento {doc_type}"), 0)
        
        # Add content
        text_content = content.get("text", "")
        for paragraph in text_content.split('\n\n'):
            if paragraph.strip():
                if paragraph.strip().startswith('#'):
                    # Header
                    header_text = paragraph.strip().lstrip('#').strip()
                    doc.add_heading(header_text, level=1)
                else:
                    # Regular paragraph
                    doc.add_paragraph(paragraph.strip())
        
        # Create output directory in project structure
        project_dir = f"projects/{project_id}"
        output_dir = os.path.join(project_dir, "generated")
        os.makedirs(output_dir, exist_ok=True)
        
        # Save document
        filename = f"{doc_type}_{doc_id}.docx"
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)
        
        return file_path
    
    def _document_needs_update(self, document: Document, new_data: Dict[str, Any]) -> bool:
        # Simple heuristic: if new data contains placeholders that this doc had, update it
        if not document.placeholders:
            return False
        
        extracted_data = new_data.get("extracted_data", {})
        return bool(extracted_data)  # For now, update if there's any new data
    
    def _update_document(self, document: Document, new_data: Dict[str, Any], db: Session) -> Dict[str, Any]:
        # Increment version
        document.version += 1
        document.status = "needs_review"
        
        # TODO: Implement intelligent document updating
        # For now, just mark it as updated
        db.commit()
        
        return {
            "document_id": document.id,
            "doc_type": document.doc_type,
            "new_version": document.version,
            "status": document.status
        }